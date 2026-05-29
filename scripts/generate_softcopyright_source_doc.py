

from __future__ import annotations



from collections import defaultdict

from dataclasses import dataclass

from datetime import datetime

from pathlib import Path

from typing import Iterable



from docx import Document

from docx.enum.section import WD_SECTION

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from docx.oxml import OxmlElement

from docx.oxml.ns import qn

from docx.shared import Cm, Pt





PROJECT_ROOT = Path(__file__).resolve().parents[1]

OUTPUT_DIR = PROJECT_ROOT / "output" / "doc"



SOFTWARE_NAME = "药途寻迹中药材全过程溯源管理系统"

SOFTWARE_VERSION = "V2.0"

DOCUMENT_DATE = datetime.now().strftime("%Y-%m-%d")



SOURCE_ROOTS = [

    PROJECT_ROOT / "后端代码" / "springboot" / "src" / "main" / "java",

    PROJECT_ROOT / "后端代码" / "springboot" / "src" / "main" / "resources",

    PROJECT_ROOT / "前端代码" / "miniprogram-5" / "src",

    PROJECT_ROOT / "前端代码" / "consumer-miniprogram" / "src",

]



ALLOWED_SUFFIXES = {

    ".java",

    ".yml",

    ".yaml",

    ".sql",

    ".html",

    ".css",

    ".js",

    ".wxml",

    ".wxss",

    ".json",

}



LINES_PER_PAGE = 50

PAGES_PER_BLOCK = 30

LINES_PER_BLOCK = LINES_PER_PAGE * PAGES_PER_BLOCK





@dataclass(frozen=True)

class FileRecord:

    file_id: int

    rel_path: str

    line_count: int





@dataclass(frozen=True)

class SourceLine:

    file_id: int

    rel_path: str

    source_line_no: int

    text: str

    segment: str





def collect_source_files() -> list[Path]:

    files: list[Path] = []

    for root in SOURCE_ROOTS:

        if not root.exists():

            continue

        files.extend(

            path

            for path in root.rglob("*")

            if path.is_file() and path.suffix.lower() in ALLOWED_SUFFIXES

        )

    return sorted(files, key=lambda path: path.relative_to(PROJECT_ROOT).as_posix())





def read_lines(path: Path) -> list[str]:

    for encoding in ("utf-8", "utf-8-sig", "gb18030"):

        try:

            text = path.read_text(encoding=encoding)

            break

        except UnicodeDecodeError:

            continue

    else:

        text = path.read_text(encoding="utf-8", errors="replace")

    lines = text.splitlines()

    return [sanitize_code_line(line) for line in lines]





def sanitize_code_line(text: str) -> str:

    text = text.expandtabs(4)

    text = text.replace("\u00a0", " ")

    text = text.replace("\x00", "")

    return text.rstrip("\n\r")





def build_catalog(files: Iterable[Path]) -> tuple[list[FileRecord], list[SourceLine]]:

    file_records: list[FileRecord] = []

    all_lines: list[SourceLine] = []



    for file_id, path in enumerate(files, start=1):

        rel_path = path.relative_to(PROJECT_ROOT).as_posix()

        lines = read_lines(path)

        file_records.append(FileRecord(file_id=file_id, rel_path=rel_path, line_count=len(lines)))

        for source_line_no, text in enumerate(lines, start=1):

            all_lines.append(

                SourceLine(

                    file_id=file_id,

                    rel_path=rel_path,

                    source_line_no=source_line_no,

                    text=text,

                    segment="",

                )

            )

    return file_records, all_lines





def select_excerpt(all_lines: list[SourceLine]) -> list[SourceLine]:

    if len(all_lines) <= LINES_PER_BLOCK * 2:

        return [

            SourceLine(

                file_id=line.file_id,

                rel_path=line.rel_path,

                source_line_no=line.source_line_no,

                text=line.text,

                segment="full",

            )

            for line in all_lines

        ]



    first_block = all_lines[:LINES_PER_BLOCK]

    last_block = all_lines[-LINES_PER_BLOCK:]

    selected = []

    for line in first_block:

        selected.append(

            SourceLine(

                file_id=line.file_id,

                rel_path=line.rel_path,

                source_line_no=line.source_line_no,

                text=line.text,

                segment="front",

            )

        )

    for line in last_block:

        selected.append(

            SourceLine(

                file_id=line.file_id,

                rel_path=line.rel_path,

                source_line_no=line.source_line_no,

                text=line.text,

                segment="back",

            )

        )

    return selected





def add_page_number(paragraph) -> None:

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()



    fld_char_begin = OxmlElement("w:fldChar")

    fld_char_begin.set(qn("w:fldCharType"), "begin")



    instr_text = OxmlElement("w:instrText")

    instr_text.set(qn("xml:space"), "preserve")

    instr_text.text = "PAGE"



    fld_char_end = OxmlElement("w:fldChar")

    fld_char_end.set(qn("w:fldCharType"), "end")



    run._r.append(fld_char_begin)

    run._r.append(instr_text)

    run._r.append(fld_char_end)





def set_code_style(document: Document) -> None:

    normal = document.styles["Normal"]

    normal.font.name = "Consolas"

    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")

    normal.font.size = Pt(8)

    fmt = normal.paragraph_format

    fmt.space_before = Pt(0)

    fmt.space_after = Pt(0)

    fmt.line_spacing = Pt(10.5)





def configure_section(section) -> None:

    section.page_width = Cm(21)

    section.page_height = Cm(29.7)

    section.top_margin = Cm(1.5)

    section.bottom_margin = Cm(1.5)

    section.left_margin = Cm(1.3)

    section.right_margin = Cm(1.3)

    section.header_distance = Cm(0.8)

    section.footer_distance = Cm(0.8)

    footer = section.footer.paragraphs[0]

    footer.clear()

    footer.style = "Normal"

    add_page_number(footer)





def build_code_document(excerpt: list[SourceLine], target_path: Path) -> None:

    document = Document()

    set_code_style(document)

    configure_section(document.sections[0])



    for index, line in enumerate(excerpt, start=1):

        paragraph = document.add_paragraph()

        paragraph.style = "Normal"

        text = f"{line.file_id:03d}:{line.source_line_no:04d}  {line.text}"

        paragraph.add_run(text)

        if index % LINES_PER_PAGE == 0 and index != len(excerpt):

            paragraph.runs[0].add_break(WD_BREAK.PAGE)



    document.save(target_path)





def build_cover_document(

    excerpt: list[SourceLine],

    file_records: list[FileRecord],

    total_source_lines: int,

    target_path: Path,

) -> None:

    document = Document()

    set_code_style(document)

    configure_section(document.sections[0])



    title = document.add_paragraph()

    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_format = title.paragraph_format

    title_format.space_before = Pt(48)

    title_format.space_after = Pt(24)

    run = title.add_run(f"{SOFTWARE_NAME}\n源代码文档")

    run.bold = True

    run.font.name = "黑体"

    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    run.font.size = Pt(18)



    info_lines = [

        f"软件版本：{SOFTWARE_VERSION}",

        f"整理日期：{DOCUMENT_DATE}",

        f"源码范围：后端 Spring Boot + 管理端小程序 + 消费端小程序",

        f"抽取规则：按文件名排序，抽取前 {PAGES_PER_BLOCK} 页与后 {PAGES_PER_BLOCK} 页，每页 {LINES_PER_PAGE} 行",

        f"源码总量：{len(file_records)} 个文件，约 {total_source_lines} 行",

        "正文行格式：文件编号:文件内行号  源码内容",

        "文件编号对照与抽取范围见同目录《软著源码抽取说明.md》",

    ]

    for line in info_lines:

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragraph.paragraph_format.space_after = Pt(6)

        paragraph.add_run(line)



    document.add_section(WD_SECTION.NEW_PAGE)

    configure_section(document.sections[-1])

    for index, line in enumerate(excerpt, start=1):

        paragraph = document.add_paragraph()

        paragraph.style = "Normal"

        paragraph.add_run(f"{line.file_id:03d}:{line.source_line_no:04d}  {line.text}")

        if index % LINES_PER_PAGE == 0 and index != len(excerpt):

            paragraph.runs[0].add_break(WD_BREAK.PAGE)



    document.save(target_path)





def build_text_excerpt(excerpt: list[SourceLine], target_path: Path) -> None:

    content = "\n".join(f"{line.file_id:03d}:{line.source_line_no:04d}  {line.text}" for line in excerpt)

    target_path.write_text(content + "\n", encoding="utf-8")





def build_manifest(

    file_records: list[FileRecord],

    excerpt: list[SourceLine],

    total_source_lines: int,

    target_path: Path,

) -> None:

    selected_file_ids = {line.file_id for line in excerpt}

    selected_files = [record for record in file_records if record.file_id in selected_file_ids]



    grouped_ranges: dict[str, list[tuple[int, str, int, int]]] = defaultdict(list)

    last_key = None

    range_start = None

    range_end = None



    for line in excerpt:

        key = (line.segment, line.file_id, line.rel_path)

        if key != last_key or range_end is None or line.source_line_no != range_end + 1:

            if last_key is not None and range_start is not None and range_end is not None:

                grouped_ranges[last_key[0]].append((last_key[1], last_key[2], range_start, range_end))

            last_key = key

            range_start = line.source_line_no

            range_end = line.source_line_no

        else:

            range_end = line.source_line_no



    if last_key is not None and range_start is not None and range_end is not None:

        grouped_ranges[last_key[0]].append((last_key[1], last_key[2], range_start, range_end))



    lines: list[str] = []

    lines.append(f"# {SOFTWARE_NAME} 软著源码抽取说明")

    lines.append("")

    lines.append(f"- 软件版本：{SOFTWARE_VERSION}")

    lines.append(f"- 生成日期：{DOCUMENT_DATE}")

    lines.append(f"- 纳入统计的源码文件数：{len(file_records)}")

    lines.append(f"- 纳入统计的源码总行数：{total_source_lines}")

    lines.append(f"- 抽取正文行数：{len(excerpt)}")

    lines.append(f"- 抽取规则：按相对路径排序，抽取前 {LINES_PER_BLOCK} 行和后 {LINES_PER_BLOCK} 行")

    lines.append(f"- 正文行格式：`文件编号:文件内行号  源码内容`")

    lines.append("")



    for label, title in (("front", "前 30 页抽取范围"), ("back", "后 30 页抽取范围"), ("full", "全文抽取范围")):

        if not grouped_ranges.get(label):

            continue

        lines.append(f"## {title}")

        lines.append("")

        for file_id, rel_path, start, end in grouped_ranges[label]:

            if start == end:

                range_text = f"{start}"

            else:

                range_text = f"{start}-{end}"

            lines.append(f"- {file_id:03d} `{rel_path}` 行 {range_text}")

        lines.append("")



    lines.append("## 文件编号对照表")

    lines.append("")

    for record in selected_files:

        lines.append(f"- {record.file_id:03d} `{record.rel_path}` 共 {record.line_count} 行")

    lines.append("")



    target_path.write_text("\n".join(lines), encoding="utf-8")





def main() -> None:

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)



    source_files = collect_source_files()

    if not source_files:

        raise SystemExit("未找到可纳入软著文档的源码文件。")



    file_records, all_lines = build_catalog(source_files)

    excerpt = select_excerpt(all_lines)

    total_source_lines = len(all_lines)



    build_code_document(excerpt, OUTPUT_DIR / "软著源码正文-60页.docx")

    build_cover_document(

        excerpt,

        file_records,

        total_source_lines,

        OUTPUT_DIR / "软著源码文档-标准版.docx",

    )

    build_text_excerpt(excerpt, OUTPUT_DIR / "软著源码正文-60页.txt")

    build_manifest(

        file_records,

        excerpt,

        total_source_lines,

        OUTPUT_DIR / "软著源码抽取说明.md",

    )



    print(f"已生成：{OUTPUT_DIR}")





if __name__ == "__main__":

    main()
