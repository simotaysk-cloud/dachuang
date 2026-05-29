

from __future__ import annotations



from datetime import date

from pathlib import Path



from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn

from docx.shared import Pt





PROJECT_ROOT = Path(__file__).resolve().parents[1]

OUTPUT_DIR = PROJECT_ROOT / "output" / "doc" / "软著申请材料"

MANIFEST_PATH = PROJECT_ROOT / "output" / "doc" / "软著源码抽取说明.md"



SOFTWARE_NAME = "药途寻迹中药材全过程溯源管理系统"

SOFTWARE_SHORT_NAME = "药途寻迹系统"

VERSION = "V1.0"

COMPLETION_DATE = "2026-04-04"

PUBLICATION_STATUS = "未发表"

RIGHT_HOLDERS = ["朱厚华", "刘少良", "张邓美颜"]

DEVELOPMENT_MODE = "合作开发"

RIGHTS_OBTAIN_MODE = "原始取得"

DOC_DATE = str(date.today())





def parse_manifest() -> tuple[str, str]:

    if not MANIFEST_PATH.exists():

        return "246", "25595"

    files = "246"

    lines = "25595"

    for raw_line in MANIFEST_PATH.read_text(encoding="utf-8").splitlines():

        line = raw_line.strip()

        if line.startswith("- 纳入统计的源码文件数："):

            files = line.removeprefix("- 纳入统计的源码文件数：")

        elif line.startswith("- 纳入统计的源码总行数："):

            lines = line.removeprefix("- 纳入统计的源码总行数：")

    return files, lines





def build_application_info(file_count: str, line_count: str) -> str:

    holders = "、".join(RIGHT_HOLDERS)

    return f"""# {SOFTWARE_NAME} 软著申请信息草稿

## 一、申请基本信息

- 软件全称：{SOFTWARE_NAME}
- 软件简称：{SOFTWARE_SHORT_NAME}
- 建议版本号：{VERSION}
- 开发完成日期：{COMPLETION_DATE}
- 发表状态：{PUBLICATION_STATUS}
- 权利取得方式：{RIGHTS_OBTAIN_MODE}
- 开发方式：{DEVELOPMENT_MODE}
- 著作权人：{holders}
- 源码文件数：{file_count}
- 源码总行数：{line_count}

## 二、申请表可直接使用的核心文案

### 1. 软件功能简介

本软件面向中药材种植、加工、质检、物流与消费查询场景，构建覆盖“种植记录、加工流转、质量检验、物流配送、终端扫码溯源、智能咨询”的全过程数字化管理平台。系统以批次为核心对象，对中药材在生产、加工、检验、运输和消费展示等关键环节的数据进行统一采集、存储、查询与展示，并支持二维码溯源、GS1 编码处理、批次谱系追踪、区块链存证接口预留以及多角色权限控制。

### 2. 主要技术特点

1. 采用 Spring Boot 构建后端服务，提供统一 REST 接口，支持身份认证、业务数据管理和追溯查询。
2. 采用微信小程序作为业务端与消费端前端载体，分别服务于生产管理人员与消费者使用。
3. 以批次编号为主线组织种植、加工、检验、物流等环节数据，实现全过程追溯和上下游谱系关联。
4. 支持二维码生成与扫码查询，可将批次追溯信息快速投放到终端展示场景。
5. 支持 GS1 编码相关字段管理和锁定机制，便于标准化标识与标签场景接入。
6. 预留区块链上链接口和交易信息查询能力，可对核心溯源数据进行可信存证扩展。
7. 集成 AI 问答接口，可结合药材批次上下文提供辅助咨询能力。

### 3. 运行环境建议

- 服务器操作系统：Windows Server 或 Linux
- 后端运行环境：JDK 17 及以上
- 数据库环境：MySQL 8.0 及以上
- 前端运行环境：微信开发者工具、微信小程序运行环境
- 网络环境：局域网或互联网 HTTPS 接入

### 4. 开发环境说明

- 后端开发语言：Java
- 前端开发语言：JavaScript、WXML、WXSS
- 后端框架：Spring Boot、Spring Data JPA
- 认证方式：JWT
- 构建工具：Maven、Node.js 生态工具链

## 三、著作权人信息待补项

以下内容我没有从你这里拿，建议你自己补，不要直接在聊天里发完整证件号：

- 朱厚华：身份证号码、联系地址、手机号码、邮箱
- 刘少良：身份证号码、联系地址、手机号码、邮箱
- 张邓美颜：身份证号码、联系地址、手机号码、邮箱

## 四、合作开发相关建议

由于本软件为合作开发，提交时建议至少准备以下材料之一：

1. 三位著作权人共同作为申请人，在申请表中共同署名。
2. 如平台或代理要求补充权属证明，准备合作开发协议或合作开发情况说明，并由三位著作权人签字确认。

## 五、目前已完成的材料

- 源代码鉴别材料：已生成前 30 页和后 30 页版本
- 文档鉴别材料：已生成《软件说明书》草稿
- 合作开发情况说明：已生成草稿

## 六、提交前最后核对

1. 软件名称、版本号、完成日期在申请表、代码材料、说明书中必须完全一致。
2. 首次登记建议使用 {VERSION}，如果你们已经在线上或比赛材料中统一写成其他版本号，需要全部同步修改。
3. 三位著作权人的姓名顺序建议固定，不要在不同材料中来回变化。
4. 如果受理方要求“60 页纯源码不带封面”，使用正文版文件提交。
"""





def build_software_manual(file_count: str, line_count: str) -> str:

    return f"""# {SOFTWARE_NAME} 软件说明书

## 1. 引言

### 1.1 编写目的

本文档用于说明 {SOFTWARE_NAME} 的建设目标、系统架构、功能模块、运行环境、业务流程与使用方式，可作为软件著作权登记的文档鉴别材料，也可作为项目后续部署、展示和维护的参考资料。

### 1.2 软件名称和版本

- 软件名称：{SOFTWARE_NAME}
- 软件简称：{SOFTWARE_SHORT_NAME}
- 软件版本：{VERSION}

### 1.3 开发背景

中药材在种植、加工、检验、运输和销售过程中涉及多个参与主体，传统纸质登记和分散式记录方式难以实现全过程追踪与快速核验。本软件围绕中药材批次管理和过程追溯场景建设数字化系统，以批次为主线整合多环节信息，实现中药材全生命周期的数据记录、查询、展示和辅助分析。

## 2. 系统概述

### 2.1 建设目标

本系统旨在建设一套面向中药材产业链的全过程溯源管理平台，实现以下目标：

1. 对种植、加工、质检、物流等关键环节信息进行结构化采集与统一管理。
2. 通过批次号和二维码实现药材来源、过程记录和流向信息的追踪查询。
3. 支持多角色协同操作，保证不同岗位用户在统一平台中完成各自业务录入与核验。
4. 通过标准化编码、数据留痕和可信存证接口，提升中药材流通过程中的真实性和可核查性。
5. 为消费者和监管场景提供便捷的终端查询与展示入口。

### 2.2 适用对象

本系统适用于中药材产业链中的农户、加工厂、质检人员、物流人员、管理员、监管人员以及终端消费者。

### 2.3 软件构成

本软件由以下三个部分构成：

1. 后端服务端：负责用户认证、业务逻辑处理、数据存储与接口提供。
2. 业务管理小程序：面向生产、加工、检验、物流和管理人员，完成业务录入、查询和维护。
3. 消费查询小程序：面向消费者提供首页展示、扫码溯源、云市集与 AI 智能咨询等功能。

## 3. 技术架构

### 3.1 总体架构

系统采用前后端分离架构。后端基于 Spring Boot 构建，负责提供统一 API 接口与业务处理能力；前端采用微信小程序实现，分为业务管理端和消费端两个应用；数据统一保存在 MySQL 数据库中；二维码、GS1 编码、批次谱系和区块链接口围绕核心批次数据展开。

### 3.2 技术选型

- 后端语言：Java
- 前端语言：JavaScript、WXML、WXSS
- 后端框架：Spring Boot 3.2.2
- 数据访问：Spring Data JPA
- 认证机制：JWT
- 数据库：MySQL 8.0 及以上
- 构建工具：Maven、Node.js 工具链

### 3.3 数据规模与代码规模

- 纳入统计的源码文件数：{file_count}
- 纳入统计的源码总行数：{line_count}
- 主要源码目录：后端代码/springboot、前端代码/miniprogram-5、前端代码/consumer-miniprogram

## 4. 功能设计

### 4.1 用户与权限管理

系统提供用户登录与身份识别功能，通过用户名、密码和令牌机制实现访问控制。系统支持管理员、农户、加工厂、物流人员、监管方等角色区分，可根据用户身份进入对应业务页面并调用对应接口。

### 4.2 批次管理

批次管理是系统核心模块。系统支持批次新增、查询、修改、删除、父子批次派生、批次上下游谱系查询、叶子批次汇总查询等功能，可将中药材的不同环节过程统一绑定到批次号之下。

### 4.3 种植记录管理

系统支持录入地块信息、农事操作、执行人员、操作时间、地理位置、现场情况说明等内容，用于记录中药材从种植到田间管理阶段的真实过程信息，提高溯源数据完整性。

### 4.4 加工记录管理

系统支持加工类型、加工产线、执行厂区、产出数量、加工说明等信息录入，可将原始批次在加工过程中派生成新的子批次，实现原料到产成品的过程留痕与流转追踪。

### 4.5 质检管理

系统支持检验结果、检验结论、检验人员、报告信息等内容管理，用于记录药材或产品的质量检测结果，便于后续追溯展示和质量核验。

### 4.6 物流与发运管理

系统支持发运单创建、物流事件记录、批量运输关系维护、物流轨迹信息展示等功能，可记录运输过程中的时间、地点、状态等数据。

### 4.7 二维码与 GS1 编码管理

系统支持生成批次二维码，便于线下扫码查看追溯信息；同时支持 GS1 编码相关字段处理、单位换算与锁定机制，为标准化标签和产品标识管理提供基础能力。

### 4.8 消费查询与追溯展示

消费端小程序提供首页、云市集、扫码追溯、个人中心等页面。消费者可通过扫码或输入批次信息查看药材来源、过程记录、检验结果、物流信息和批次状态，从而了解药材的全过程信息。

### 4.9 AI 辅助咨询

系统提供 AI 对话接口，可结合当前扫码药材或批次上下文，为用户提供药材介绍、使用建议、注意事项和配伍参考等辅助咨询功能。

### 4.10 区块链存证扩展

系统预留区块链存证接口，支持记录批次数据摘要、交易哈希、链上访问链接等信息，为核心追溯数据的可信留痕和后续链上验证提供扩展空间。

## 5. 前端页面说明

### 5.1 业务管理端页面

业务管理小程序主要包含以下页面：

1. 登录页
2. 首页
3. 批次列表页
4. 批次新增与外部批次接入页
5. 批次谱系与追溯详情页
6. 种植管理页与种植表单页
7. 加工管理页与加工表单页
8. 质检管理页与质检表单页
9. 物流管理页与发运表单页
10. 二维码页与终端二维码页
11. 用户管理页
12. 数据看板页
13. 个人中心页

### 5.2 消费查询端页面

消费端小程序主要包含以下页面：

1. 首页
2. 溯源查询页
3. 云市集页
4. AI 智问页
5. 个人中心页

## 6. 核心业务流程

### 6.1 中药材全过程溯源流程

1. 管理员或业务人员创建原始批次。
2. 种植环节录入地块、农事操作与现场信息。
3. 加工环节根据原始批次派生产成品或半成品批次。
4. 质检人员录入检验结果和报告信息。
5. 物流人员创建发运记录并录入运输过程事件。
6. 系统生成二维码并绑定到批次追溯页面。
7. 消费者扫码后查看该批次的全过程追溯信息。

### 6.2 用户登录流程

1. 用户输入账号和密码。
2. 后端校验用户信息并生成认证令牌。
3. 前端保存令牌并根据角色进入不同业务页面。
4. 后续接口访问时携带令牌完成身份识别。

## 7. 运行与部署说明

### 7.1 服务器运行环境

- 操作系统：Windows Server、Ubuntu、CentOS 或其他 Linux 发行版
- Java 环境：JDK 17 及以上
- 数据库环境：MySQL 8.0 及以上
- 浏览器环境：用于管理后台静态页面展示和调试

### 7.2 小程序运行环境

- 微信开发者工具
- 微信客户端
- HTTPS 后端接口地址

### 7.3 部署步骤概述

1. 创建数据库并执行表结构迁移。
2. 配置后端数据库连接、端口和相关业务参数。
3. 启动 Spring Boot 后端服务。
4. 在微信开发者工具中导入两个小程序项目。
5. 配置接口基础地址并进行联调。

## 8. 数据安全与可维护性

系统采用统一接口层与业务层分离设计，使用令牌认证机制对用户访问进行控制。系统通过结构化数据表管理批次、种植、加工、检验、物流等业务数据，便于维护和扩展；通过统一 API 输出和异常处理机制，提高系统稳定性与可维护性。

## 9. 结论

{SOFTWARE_NAME} 围绕中药材全过程追溯场景，构建了覆盖生产、加工、检测、物流和消费查询的数字化管理平台。系统具备批次管理、追溯查询、二维码生成、标准化编码处理、区块链扩展与智能咨询等功能，能够满足中药材产业链多角色协同和全过程信息追踪的实际需求。
"""





def build_cooperation_statement() -> str:

    holders = "、".join(RIGHT_HOLDERS)

    return f"""# 合作开发情况说明（草稿）

兹说明：{SOFTWARE_NAME}（版本号：{VERSION}）由 {holders} 共同参与开发完成，开发完成日期为 {COMPLETION_DATE}。

本软件属于合作开发形成的作品，各开发者共同参与了系统需求分析、架构设计、程序编码、前端页面开发、接口联调、测试优化及文档整理等工作，对软件整体功能实现均作出了实质性贡献。

各方一致确认：

1. 本软件著作权由 {holders} 共同享有。
2. 本软件系原始开发取得，不存在侵犯第三方著作权或其他知识产权的情形。
3. 本软件目前未公开发表。
4. 各方同意以共同著作权人的身份申请计算机软件著作权登记。

本说明仅作为软件著作权登记材料准备使用，正式提交时可根据受理要求补充签字、身份证号码、联系方式及日期信息。

签字栏：

- 朱厚华：________________
- 刘少良：________________
- 张邓美颜：________________

日期：________________
"""





def build_checklist() -> str:

    return f"""# 软著提交前待补清单

## 你已经有的

- 三位著作权人姓名
- 软件名称
- 完成日期
- 未发表状态
- 源代码鉴别材料
- 软件说明书草稿

## 还需要你自己补的

1. 三位著作权人的身份证正反面扫描件或照片
2. 三位著作权人的身份证号码、联系地址、手机号、邮箱
3. 如平台要求，三位著作权人的签字页
4. 如平台或代理要求，合作开发协议或合作开发情况说明签字版
5. 最终确认版本号是否采用 {VERSION}

## 我建议的最终提交口径

- 软件名称：{SOFTWARE_NAME}
- 版本号：{VERSION}
- 开发完成日期：{COMPLETION_DATE}
- 开发方式：{DEVELOPMENT_MODE}
- 权利取得方式：{RIGHTS_OBTAIN_MODE}
- 发表状态：{PUBLICATION_STATUS}
- 文档鉴别材料：软件说明书
"""





def set_default_style(document: Document) -> None:

    style = document.styles["Normal"]

    style.font.name = "Times New Roman"

    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    style.font.size = Pt(11)





def add_markdown_to_docx(document: Document, markdown_text: str) -> None:

    for raw_line in markdown_text.splitlines():

        line = raw_line.rstrip()

        if not line:

            document.add_paragraph("")

            continue

        if line.startswith("# "):

            document.add_heading(line[2:], level=1)

        elif line.startswith("## "):

            document.add_heading(line[3:], level=2)

        elif line.startswith("### "):

            document.add_heading(line[4:], level=3)

        elif line.startswith("- "):

            document.add_paragraph(line[2:], style="List Bullet")

        elif line[:2].isdigit() and line[1:3] == ". ":

            document.add_paragraph(line[3:], style="List Number")

        else:

            document.add_paragraph(line)





def write_docx(path: Path, title: str, markdown_text: str) -> None:

    document = Document()

    set_default_style(document)

    heading = document.add_paragraph()

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = heading.add_run(title)

    run.bold = True

    run.font.name = "黑体"

    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    run.font.size = Pt(16)

    document.add_paragraph("")

    add_markdown_to_docx(document, markdown_text)

    document.save(path)





def main() -> None:

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    file_count, line_count = parse_manifest()



    application_info = build_application_info(file_count, line_count)

    software_manual = build_software_manual(file_count, line_count)

    cooperation_statement = build_cooperation_statement()

    checklist = build_checklist()



    files = {

        "软著申请信息草稿.md": application_info,

        "软件说明书-软著版.md": software_manual,

        "合作开发情况说明-草稿.md": cooperation_statement,

        "提交前待补清单.md": checklist,

    }



    for filename, content in files.items():

        (OUTPUT_DIR / filename).write_text(content, encoding="utf-8")



    write_docx(OUTPUT_DIR / "软著申请信息草稿.docx", "软著申请信息草稿", application_info)

    write_docx(OUTPUT_DIR / "软件说明书-软著版.docx", "软件说明书", software_manual)

    write_docx(OUTPUT_DIR / "合作开发情况说明-草稿.docx", "合作开发情况说明", cooperation_statement)



    print(f"已生成：{OUTPUT_DIR}")





if __name__ == "__main__":

    main()
